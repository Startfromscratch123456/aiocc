#!/usr/bin/python
# -*- coding: UTF-8 -*-
import os
import re
import datetime
import numpy as np
import matplotlib.pyplot as plt
from .FileUtils import *
from .PrintUtil import print_fio_document
import docx
from docx import *
from docx.shared import *
from FileUtils import *
from PrintUtil import print_fio_document
from FioDocument import *
from FioStatistic import *


class TestReport:
    def __init__(self):
        self.fontsize = {'three': 16, 'five': 10.5}
        self.fio_document_statistic = FioDocumentStatistic()
        self.fio_document_statistic.load_data('data')

    def __add_normal_str(self, paragraph, str):
        run = paragraph.add_run(str)
        font = run.font
        font.size = Pt(self.fontsize['five'])
        font.name = '微软雅黑 Light'

    def get_report(self):
        document = Document()
        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()
        count = [0, 0, 0, 0]
        count[0] = count[1] = count[2] = count[3] = 0  # chapter
        document.add_heading('sscdt 测试详细报告', 0)
        now = datetime.datetime.now()
        paragraph = document.add_paragraph()
        run = paragraph.add_run('时间 ' + str(now.strftime("%Y-%m-%d %H:%M:%S")))
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

        document.add_page_break()

        count[0] += 1
        count[1] = 0

        '''
                                                page 1
        '''
        paragraph = document.add_heading('%d 说明' % (count[0]), 1)
        count[1] += 1
        paragraph = document.add_heading('  %d.%d 测试环境' % (count[0], count[1]), 2)
        paragraph = document.add_paragraph()
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '项目'
        hdr_cells[1].text = '详情'
        recordset = [{'property': 'Lustre nrs crrn', 'detail': 'lustre-2.8.0 [Linux内核：linux-3.10.0-327.3.1.el7]'},
                     {'property': 'lustre nrs sscdt', 'detail': 'lustre-2.8.0 [Linux内核：linux-3.10.0-327.3.1.el7]'},
                     {'property': 'fio（测试工具）', 'detail': 'fio-2.7'}]
        for item in recordset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['property'])
            row_cells[1].text = str(item['detail'])
        table.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'LightShading-Accent1'

        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '参数名称'
        hdr_cells[1].text = '参数'
        recordset = [{'property': 'CPU', 'detail': 'Intel(R) Xeon(R) CPU E5-2620 0 @ 2.00GHz 2x '},
                     {'property': '内存', 'detail': '16GB DDR3 RAM '},
                     {'property': '硬盘',
                      'detail': '2x SAS Disks (15000RPM, 146GB) configured RAID1,6x SAS Disks (15000RPM, 146GB)configured RAID5 '},
                     {'property': '网卡', 'detail': 'Mellanox InfiniBand QDR 40Gb/s NIC'}]

        for item in recordset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['property'])
            row_cells[1].text = str(item['detail'])
        table.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        table.style = 'LightShading-Accent1'

        count[1] += 1
        paragraph = document.add_heading('  %d.%d 说明' % (count[0], count[1]), 2)
        paragraph = document.add_paragraph()
        self.__add_normal_str(paragraph,
                              "   背景：Lustre Server-Side-IO Coordination")
        paragraph = document.add_paragraph()
        self.__add_normal_str(paragraph, "    结论：sscdt对Lusre QoS有较大的提升。")

        count[1] += 1
        paragraph = document.add_heading('  %d.%d 测试参数' % (count[0], count[1]), 2)
        paragraph = document.add_paragraph()
        self.__add_normal_str(paragraph, "    一共11节点：mgs 1,oss 3 ,client2 ")

        count[0] += 1
        paragraph = document.add_heading('  %d 测试数据对比及分析' % (count[0]), 1)
        '''
                                                page 2
        '''
        count[1] = 0
        count[1] += 1
        self.write_singledoc_to_file(document, count, 'read', 'read(顺序读)')
        count[1] += 1
        self.write_singledoc_to_file(document, count, 'randread', 'randread(随机读)')
        count[1] += 1
        self.write_singledoc_to_file(document, count, 'write', 'write(顺序写)')
        count[1] += 1
        self.write_singledoc_to_file(document, count, 'randwrite', 'randwrite(随机写)')
        count[1] += 1
        self.write_singledoc_to_file(document, count, 'readwrite', 'readwrite(顺序混合读写)')
        count[1] += 1
        self.write_singledoc_to_file(document, count, 'randrw', 'randrw(随机混合读写)')
        document.save("report" + os.sep + str(now.strftime("reportV%Y%m%d%H%M%S")) + '.docx')

    def write_singledoc_to_file(self, document, count, io_pattern, title):

        paragraph = document.add_heading('  %d.%d %s' % (count[0], count[1], title), 2)
        count[2] = 0
        count[2] += 1
        paragraph = document.add_heading('  %d.%d.%d bandwidth' % (count[0], count[1], count[2]), 3)
        document.add_paragraph()

        paragraph = document.add_heading('  %d.%d.%d io(数据读取总量)' % (count[0], count[1], count[2]), 3)
        document.add_paragraph()
        if (io_pattern == 'readwrite' or io_pattern == 'randrw'):
            self.fio_document_statistic.get_rwcomparison_io(io_pattern, 'io')
        else:
            self.fio_document_statistic.get_single_rwcomparison_io(io_pattern, 'io')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'io' + '.png'))

        count[2] += 1
        paragraph = document.add_heading('  %d.%d.%d iops' % (count[0], count[1], count[2]), 3)
        document.add_paragraph()
        if (io_pattern == 'readwrite' or io_pattern == 'randrw'):
            self.fio_document_statistic.get_rwcomparison_iops(io_pattern, 'iops')
        else:
            self.fio_document_statistic.get_single_rwcomparison_iops(io_pattern, 'iops')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'iops' + '.png'))

        count[2] += 1
        paragraph = document.add_heading('  %d.%d.%d  IO完成时延（最值、标准差）' % (count[0], count[1], count[2]), 3)
        count[3] = 0
        count[3] += 1
        paragraph = document.add_heading('  %d）最小时延' % count[3], 4)
        document.add_paragraph()
        self.fio_document_statistic.get_single_rwcomparison_clat(io_pattern, 'min')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + '_clat_' + 'min' + '.png'))

        count[3] += 1
        paragraph = document.add_heading('  %d）最大时延' % count[3], 4)
        document.add_paragraph()
        self.fio_document_statistic.get_single_rwcomparison_clat(io_pattern, 'max', 'complete latency(k us)')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + '_clat_' + 'max' + '.png'))

        count[3] += 1
        paragraph = document.add_heading('  %d）平均时延' % count[3], 4)
        document.add_paragraph()
        self.fio_document_statistic.get_single_rwcomparison_clat(io_pattern, 'avg')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + '_clat_' + 'avg' + '.png'))

        count[3] += 1
        paragraph = document.add_heading('  %d）时延标准差' % count[3], 4)
        document.add_paragraph()
        self.fio_document_statistic.get_single_rwcomparison_clat(io_pattern, 'stdev')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + '_clat_' + 'stdev' + '.png'))

        count[2] += 1
        count[3] = 1
        paragraph = document.add_heading('  %d.%d.%d 完成时延详细情况:clat延迟百分位数' % (count[0], count[1], count[2]), 3)
        paragraph = document.add_paragraph()
        self.__add_normal_str(paragraph, "    说明：首先明确百分位意义，sscdt在1.00为14us，表示所有IO中1.00%的IO都是在14us中完成的。")
        bs = 1;
        while bs < (2048 + 1):
            limit = 12
            if (bs > 64):
                limit = 8
            elif (bs > 16):
                limit = 10
            paragraph = document.add_heading('  %d）clat延迟百分位数(%d k)' % (count[3], bs), 4)
            document.add_paragraph("X%以前")
            if (io_pattern == 'readwrite' or io_pattern == 'randrw'):
                self.fio_document_statistic.get_rwcomparison_clatA(io_pattern, bs, limit)
                document.add_picture(
                    str('images' + os.sep + io_pattern + os.sep + 'clat_rwA' + str(bs) + '.png'))

                self.fio_document_statistic.get_rwcomparison_clatB(io_pattern, bs, limit)
                document.add_paragraph("X%以后")
                document.add_picture(
                    str('images' + os.sep + io_pattern + os.sep + 'clat_rwB' + str(bs) + '.png'))
            else:
                self.fio_document_statistic.get_single_rwcomparison_clatA(io_pattern, bs, limit)
                document.add_picture(
                    str('images' + os.sep + io_pattern + os.sep + 'clat_percentilesA' + str(bs) + '.png'))

                self.fio_document_statistic.get_single_rwcomparison_clatB(io_pattern, bs, limit)
                document.add_paragraph("X%以后")
                document.add_picture(
                    str('images' + os.sep + io_pattern + os.sep + 'clat_percentilesB' + str(bs) + '.png'))

            count[3] += 1
            bs = (bs * 2)

        count[2] += 1;
        count[3] = 0;

        paragraph = document.add_heading('  %d.%d.%d 聚合IO(aggregate bandwidth)' % (count[0], count[1], count[2]), 3)

        document.add_paragraph()
        count[3] += 1
        paragraph = document.add_heading('  %d）聚合IO最大带宽' % count[3], 4)
        self.fio_document_statistic.get_single_rwcomparison_aggregate(io_pattern, 'max',
                                                                      'max of aggregate bandwidth(KB/S)')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'aggrebwmax' + '.png'))
        count[3] += 1

        document.add_paragraph()
        paragraph = document.add_heading('  %d）聚合IO的平均带宽' % count[3], 4)
        self.fio_document_statistic.get_single_rwcomparison_aggregate(io_pattern, 'avg',
                                                                      'avg of aggregate bandwidth(KB/S)')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'aggrebwavg' + '.png'))

        count[3] += 1
        document.add_paragraph()
        paragraph = document.add_heading('  %d）聚合IO的标准差(Standard Deviation)' % count[3], 4)
        self.fio_document_statistic.get_single_rwcomparison_aggregate(io_pattern, 'stdev',
                                                                      'stdev average of aggregate bandwidth(KB/S)')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'aggrebwstdev' + '.png'))

        count[3] += 1
        document.add_paragraph()
        paragraph = document.add_heading('  %d）聚合IO占的百分比' % count[3], 4)
        self.fio_document_statistic.get_single_rwcomparison_aggregate(io_pattern, 'per',
                                                                      'per of aggregate bandwidth( %)')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'aggrebwper' + '.png'))
        # count[3] += 1

        count[2] += 1
        count[3] = 0

        paragraph = document.add_heading('  %d.%d.%d CPU使用情况' % (count[0], count[1], count[2]), 3)
        count[3] += 1
        document.add_paragraph()
        paragraph = document.add_heading('  %d）用户占用时间百分比' % count[3], 4)
        self.fio_document_statistic.get_single_rwcomparison_cpu(io_pattern, 'usr',
                                                                'CPU usage statistics - user time( %)')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'cpuusr' + '.png'))

        count[3] += 1
        document.add_paragraph()
        paragraph = document.add_heading('  %d）系统占用时间百分比' % count[3], 4)
        self.fio_document_statistic.get_single_rwcomparison_cpu(io_pattern, 'sys',
                                                                'CPU usage statistics - sys time( %)')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'cpusys' + '.png'))

        count[3] += 1
        document.add_paragraph()
        paragraph = document.add_heading('  %d）CPU上下文切换次数' % count[3], 4)
        self.fio_document_statistic.get_single_rwcomparison_cpu(io_pattern, 'ctx', 'number of context switches')
        document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'cpuctx' + '.png'))
        document.add_page_break()

        count[2] += 1
        count[3] = 0
        count[3] += 1
        paragraph = document.add_heading('  %d.%d.%d 任务完成时延分布' % (count[0], count[1], count[2]), 3)
        paragraph = document.add_paragraph()
        self.__add_normal_str(paragraph, "    说明：(a,b]表示在时间段a-b(包括b)内完成任务所占总任务百分比")
        bs = 1;
        while bs < (2048 + 1):
            paragraph = document.add_heading('  %d）Distribution of I/O completion latencies(%d k)' % (count[3], bs), 4)
            document.add_paragraph("微秒级")
            self.fio_document_statistic.get_all_rwcomparison_clatA(io_pattern, bs)
            document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'rwclatus' + str(bs) + '.png'))
            document.add_page_break()
            self.fio_document_statistic.get_all_rwcomparison_clatB(io_pattern, bs)
            document.add_paragraph("毫秒级")
            document.add_picture(str('images' + os.sep + io_pattern + os.sep + 'rwclatms' + str(bs) + '.png'))
            count[3] += 1
            bs = (bs * 2)
